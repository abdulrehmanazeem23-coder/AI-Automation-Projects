# utils/docx_exporter.py
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_proposal_to_docx(rfp_data: dict, sections: list, compliance: dict) -> bytes:
    """
    Build a professional Word document from the AI-drafted proposal sections.
    Returns bytes that Streamlit can offer as a download.
    """
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover / Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("PROPOSAL RESPONSE")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    doc.add_paragraph()

    project_para = doc.add_paragraph()
    project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj_run = project_para.add_run(rfp_data.get("title", "Untitled Project"))
    proj_run.bold = True
    proj_run.font.size = Pt(16)

    doc.add_paragraph()

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(
        f"Submitted to: {rfp_data.get('client', '—')}\n"
        f"Sector: {rfp_data.get('sector', '—')}  |  "
        f"Budget: {rfp_data.get('budget') or '—'}  |  "
        f"Deadline: {rfp_data.get('deadline') or '—'}"
    ).font.size = Pt(11)

    doc.add_page_break()

    # ── Executive Summary ──
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(rfp_data.get("summary", "No summary available."))
    doc.add_paragraph()

    # ── Compliance Overview ──
    doc.add_heading("Compliance Overview", level=1)
    comp_intro = doc.add_paragraph(
        f"Our organization meets {compliance['passed']} out of {compliance['total']} "
        f"mandatory requirements ({compliance['compliance_rate']}% compliance rate)."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Requirement"
    hdr[1].text = "Status"
    hdr[2].text = "Evidence"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for item in compliance["items"]:
        row = table.add_row().cells
        row[0].text = item["requirement"][:120]
        row[1].text = "✅ Pass" if "Pass" in item["status"] else "❌ Gap"
        row[2].text = item["best_match"][:100]

    doc.add_paragraph()
    doc.add_page_break()

    # ── Proposal Sections ──
    doc.add_heading("Technical Proposal", level=1)
    doc.add_paragraph(
        "The following sections address each question and requirement outlined in the RFP. "
        "All responses are supported by documented past project experience."
    )
    doc.add_paragraph()

    for i, sec in enumerate(sections, 1):
        doc.add_heading(f"Section {i}: {sec['question'][:80]}", level=2)
        doc.add_paragraph(sec["drafted_text"])

        # Evidence footnote
        if sec.get("evidence_used"):
            evid_para = doc.add_paragraph()
            evid_run = evid_para.add_run(
                "Supporting Evidence: " +
                ", ".join([f"[{m['cap_id']}] {m['domain']}" for m in sec["evidence_used"]])
            )
            evid_run.italic = True
            evid_run.font.size = Pt(9)
            evid_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

    # ── Save to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()